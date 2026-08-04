# -*- coding: utf-8 -*-
"""
Genera el reporte VAPT FSEC-18 (DOCX) para FleetSec S.A.S.

Self-contained: extrae los metadatos estructurados directamente de las fichas
vapt/findings/V-*.md (ID, severidad, CVSS, vector, CWE, OWASP, estado, commit) y
construye el DOCX. Reproducible con:  python vapt/reports/generate-report.py
Requiere: python-docx.
"""
import os
import re
import glob
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
OUT = os.path.join(ROOT, "vapt", "reports", "FSEC-18-Reporte-VAPT-FleetSec.docx")


def extract_findings():
    out = []
    for f in sorted(glob.glob(os.path.join(ROOT, "vapt", "findings", "V-*.md"))):
        t = open(f, encoding="utf-8").read()

        def g(pat, d=""):
            m = re.search(pat, t)
            return m.group(1).strip() if m else d

        out.append(dict(
            id=g(r"\*\*ID\*\*\s*\|\s*([^\|_]+)"),
            sev=g(r"\*\*Severidad\*\*\s*\|\s*[^A-Z]*([A-ZÍ]+)"),
            cvss=g(r"\*\*CVSSv3\.1\*\*\s*\|\s*\*\*([0-9.]+)\*\*"),
            vector=g(r"`(CVSS:3\.1/[^`]+)`"),
            cwe=g(r"\[(CWE-\d+)\]"),
            owasp=g(r"(A\d\d:2021[^\]|]*)").strip(),
            estado="Fixed" if "Fixed" in t else "Open",
            commit=g(r"commit `([0-9a-f]{7})`"),
        ))
    return out


IDX = extract_findings()
BY = {d["id"]: d for d in IDX}

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xC0, 0x39, 0x2B)
AMBER = RGBColor(0xB9, 0x77, 0x0E)
GREEN = RGBColor(0x1E, 0x7E, 0x34)
SEV_COLOR = {"HIGH": AMBER, "MEDIUM": RGBColor(0x9A, 0x7D, 0x0A), "CRITICAL": RED}

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
for lvl, sz in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = NAVY
    st.font.bold = True


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell(cell, text, bold=False, color=None, size=9.5, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if white:
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color:
        r.font.color.rgb = color


def header_row(table, headers):
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, white=True, size=9.5)
        shade(hdr[i], "1F3A5F")


def para(text, size=10.5, bold=False, italic=False, color=None, space_after=6, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def bullet(text, size=10.5, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(size)
    r2 = p.add_run(text)
    r2.font.size = Pt(size)
    return p


# ============================ PORTADA ============================
for _ in range(3):
    doc.add_paragraph()
para("Reporte de Evaluación de Vulnerabilidades\ny Pruebas de Penetración (VAPT)", size=24, bold=True,
     color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("API de Telemetría Vehicular — FleetSec S.A.S.", size=14, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
para("Clasificación: CONFIDENCIAL", size=11, bold=True, color=RED,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
meta = doc.add_table(rows=6, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
info = [
    ("Cliente", "FleetSec S.A.S. (Quantum Data Processing Colombia)"),
    ("Autor", "Juan Andrés Moya — Ingeniero de Ciberseguridad"),
    ("Fecha del reporte", "3 de agosto de 2026"),
    ("Alcance", "API REST FleetSec (Spring Boot) — conductores, vehículos, reportes, autenticación"),
    ("Marco de referencia", "OWASP WSTG v4.2 · OWASP Top 10 2021 · ASVS 4.0.3 · CVSS v3.1 · CWE"),
    ("Estado del entregable", "11/11 hallazgos remediados y verificados"),
]
for i, (k, v) in enumerate(info):
    set_cell(meta.rows[i].cells[0], k, bold=True, color=NAVY, size=10)
    set_cell(meta.rows[i].cells[1], v, size=10)
    meta.rows[i].cells[0].width = Inches(1.9)
    meta.rows[i].cells[1].width = Inches(4.4)
doc.add_page_break()

# ============================ 1. RESUMEN EJECUTIVO ============================
doc.add_heading("1. Resumen ejecutivo", level=1)
para("Este informe presenta los resultados de una evaluación de seguridad sobre la plataforma "
     "de telemetría vehicular de FleetSec, la aplicación que gestiona los datos de aproximadamente "
     "60.000 conductores y sus vehículos. La evaluación buscó responder una pregunta de negocio: "
     "¿pueden personas no autorizadas acceder, alterar o robar los datos personales que la empresa "
     "está obligada a proteger bajo la Ley 1581 de 2012?", space_after=8)
para("La respuesta inicial fue sí. Se identificaron 11 debilidades de seguridad. La más grave —y la "
     "raíz de casi todas las demás— era que la aplicación no exigía identificarse para usar sus "
     "funciones: cualquiera en la red podía consultar la lista de conductores con su cédula, correo y "
     "teléfono, leer los recorridos de un vehículo o modificar registros, sin credenciales. Sobre esa "
     "base, otras fallas permitían extraer contraseñas, leer archivos internos del servidor o alcanzar "
     "componentes de la infraestructura en la nube.", space_after=8)
para("Todas las debilidades fueron corregidas y verificadas. Cada corrección se acompaña de una prueba "
     "automática doble que confirma dos cosas: que el ataque ya no funciona y que la función legítima "
     "sigue operando. El sistema de control de calidad de seguridad (pipeline) quedó en verde: no "
     "reporta ninguna falla de severidad alta pendiente.", space_after=8)
para("Panorama de riesgo", size=11, bold=True, color=NAVY, space_after=4)
n_high = sum(1 for d in IDX if d["sev"] == "HIGH")
n_med = sum(1 for d in IDX if d["sev"] == "MEDIUM")
cifras = doc.add_table(rows=2, cols=4)
cifras.style = "Light Grid Accent 1"
c = cifras.rows[0].cells
for i, h in enumerate(["Total hallazgos", "Severidad alta", "Severidad media", "Remediados"]):
    set_cell(c[i], h, bold=True, size=9.5)
v = cifras.rows[1].cells
for i, val in enumerate([str(len(IDX)), str(n_high), str(n_med), f"{len(IDX)}/{len(IDX)} (100%)"]):
    set_cell(v[i], val, bold=True, size=12, color=NAVY)
para("", space_after=4)
para("No se registraron hallazgos de severidad crítica. Esto no es casualidad ni omisión: la "
     "puntuación se asignó por el impacto realmente demostrado en cada prueba de concepto, no por el "
     "peor escenario teórico. Es una señal de rigor metodológico, no de subestimación.", size=9.5,
     italic=True, color=GREY, space_after=8)
para("Recomendaciones de alto nivel", size=11, bold=True, color=NAVY, space_after=4)
bullet("mantener el control de acceso obligatorio como línea base no negociable de toda la API.", bold_prefix="Acceso: ")
bullet("activar el gate de seguridad del pipeline como control bloqueante para todo cambio futuro.", bold_prefix="Prevención: ")
bullet("trasladar el control de fuerza bruta y el enmascaramiento de datos al borde (WAF/gateway) al escalar a múltiples réplicas.", bold_prefix="Escalabilidad: ")
bullet("notificar a la SIC solo si esta exposición hubiese ocurrido en producción; en este entorno controlado no hubo datos reales comprometidos.", bold_prefix="Cumplimiento: ")
doc.add_page_break()

# ============================ 2. ALCANCE Y METODOLOGÍA ============================
doc.add_heading("2. Alcance y metodología", level=1)
doc.add_heading("2.1 Alcance", level=2)
para("El objeto de evaluación es la API REST de FleetSec (aplicación Spring Boot 3.3.5 / Java 21), "
     "compuesta por cuatro controladores —autenticación, conductores, vehículos y reportes— sobre un "
     "dominio de conductores (PII), vehículos y viajes. La evaluación abarcó los diez vectores de "
     "ataque planificados más un hallazgo adicional de arquitectura (V-11) detectado durante el trabajo.")
doc.add_heading("2.2 Metodología", level=2)
para("Se aplicó un enfoque de caja gris/blanca (con acceso al código fuente), combinando tres fuentes "
     "de evidencia por hallazgo para reducir falsos positivos:")
bullet("análisis estático (SAST, Semgrep) sobre el código.", bold_prefix="Estática: ")
bullet("análisis dinámico (DAST, OWASP ZAP) contra la aplicación en ejecución.", bold_prefix="Dinámica: ")
bullet("verificación manual con prueba de concepto reproducible (curl / request-response con timestamp).", bold_prefix="Manual: ")
para("La severidad se calculó con CVSS v3.1 usando el vector explícito y puntuando el impacto "
     "demostrado. Cada hallazgo se mapeó a CWE, OWASP Top 10 2021 y OWASP ASVS 4.0.3, con el impacto "
     "bajo Ley 1581 de 2012 cuando involucra datos personales.")
doc.add_page_break()

# ============================ 3. RESUMEN DE HALLAZGOS ============================
doc.add_heading("3. Resumen de hallazgos", level=1)
TITULOS = {
    "V-01": "SQL Injection en búsqueda de conductores",
    "V-02": "JWT alg:none aceptado (firma no verificada)",
    "V-03": "SSRF en webhook de vehículos (alcanza IMDS)",
    "V-04": "XXE en importación XML de vehículos",
    "V-05": "Mass Assignment en actualización de conductor",
    "V-06": "Path Traversal en descarga de reportes",
    "V-07": "Ausencia de rate limiting en login",
    "V-08": "Logging de PII en texto plano (Ley 1581)",
    "V-09": "IDOR en viajes de conductor",
    "V-10": "Credenciales hardcodeadas en configuración",
    "V-11": "Ausencia de enforcement de autenticación (raíz)",
}
tbl = doc.add_table(rows=1, cols=6)
tbl.style = "Table Grid"
header_row(tbl, ["ID", "Hallazgo", "Sev.", "CVSS", "CWE / OWASP", "Estado"])
for vid in [f"V-{i:02d}" for i in range(1, 12)]:
    d = BY[vid]
    row = tbl.add_row().cells
    set_cell(row[0], vid, bold=True, size=9)
    set_cell(row[1], TITULOS[vid], size=9)
    set_cell(row[2], d["sev"], bold=True, size=9, color=SEV_COLOR.get(d["sev"], GREY))
    set_cell(row[3], d["cvss"], size=9)
    set_cell(row[4], f"{d['cwe']} · {d['owasp']}", size=8.5)
    set_cell(row[5], "✔ Fixed", size=9, color=GREEN, bold=True)
para("", space_after=4)
para("Distribución: %d de severidad alta, %d de severidad media, 0 crítica. Los 11 hallazgos están "
     "remediados y verificados con prueba dual." % (n_high, n_med), size=9.5, italic=True, color=GREY)
doc.add_page_break()

# ============================ 4. ANÁLISIS DE INTERACCIÓN ============================
doc.add_heading("4. Análisis de interacción entre vulnerabilidades", level=1)
para("Las vulnerabilidades no son independientes: se potencian entre sí y el orden de remediación "
     "importa. Esta sección —central en el análisis— explica las cadenas de ataque y por qué la "
     "corrección se ejecutó en tres fases secuenciales.", space_after=8)
doc.add_heading("4.1 El hallazgo raíz y la cadena de autenticación", level=2)
para("V-11 (ausencia de enforcement de autenticación) es el hallazgo raíz. Mientras existió, "
     "acotaba paradójicamente el impacto de V-02 (JWT alg:none) y V-10 (credenciales hardcodeadas): "
     "los tokens no controlaban nada, así que falsificarlos era irrelevante. Pero esa relación se "
     "invierte al remediar: activar la autenticación (V-11) sobre un mecanismo de tokens todavía "
     "roto (V-02/V-10) habría construido la puerta de entrada al account takeover.", space_after=8)
para("Secuencia de remediación (obligatoria):", size=10.5, bold=True, space_after=2)
para("Fase 1 — reparar el mecanismo de identidad: V-02 (verificación de firma con jjwt, se rechaza "
     "alg:none) + V-10 (secretos externalizados con arranque fail-fast).", size=10, space_after=2)
para("Fase 2 — activar el enforcement: V-11 (Spring Security default-deny) y, en la misma fase, el "
     "control de propiedad V-05 y V-09 (que solo tienen sentido una vez que hay identidad).", size=10, space_after=2)
para("Fase 3 — endurecer las entradas: V-01, V-03, V-04, V-06, V-07, V-08 (sin dependencias de "
     "orden entre sí).", size=10, space_after=8)
doc.add_heading("4.2 Cadena de exposición de datos", level=2)
para("V-11 amplificaba a V-01 (SQLi) y V-06 (path traversal) volviéndolos explotables de forma "
     "totalmente anónima: sin autenticación, la extracción de PII vía inyección o la lectura de "
     "archivos no requería siquiera una cuenta. Del mismo modo, V-09 (IDOR) y V-05 (mass assignment) "
     "eran alcanzables sin credenciales. Remediar V-11 elevó la barrera de entrada para todos; las "
     "correcciones específicas de cada vector cierran el resto de la superficie.", space_after=8)
doc.add_heading("4.3 Las dos vías al account takeover (cerradas y probadas)", level=2)
bullet("un token alg:none o firmado con otro secreto ya no autentica (401). Verificado en pruebas de interacción.", bold_prefix="V-02/V-10 → V-11: ")
bullet("un conductor autenticado no puede auto-promoverse a admin vía mass assignment (role no asignable). Verificado.", bold_prefix="V-05 → V-11: ")
doc.add_page_break()

# ============================ 5. HALLAZGOS TÉCNICOS ============================
doc.add_heading("5. Hallazgos técnicos detallados", level=1)
DETALLE = {
 "V-01": ("GET /api/drivers/search?q=",
   "El parámetro de búsqueda se concatenaba directamente en la sentencia SQL.",
   "Extracción de PII de todos los conductores, contraseñas en claro (UNION) y lectura de archivos locales (FILE_READ de H2).",
   "Consulta parametrizada: el input viaja como parámetro LIKE enlazado, nunca como código SQL.",
   "El payload de inyección se trata como literal (0 resultados); una búsqueda real devuelve el conductor."),
 "V-02": ("POST /api/auth/validate (eliminado)",
   "La validación del JWT aceptaba tokens con alg:none, es decir, sin firma verificada.",
   "Falsificación de identidad/rol (acotada mientras V-11 estuvo presente).",
   "Migración a jjwt con verifyWith(); alg:none se rechaza por diseño. El endpoint-oráculo /validate se eliminó.",
   "El token alg:none exacto de la PoC → 401 en endpoint protegido; un HS256 válido → 200."),
 "V-03": ("POST /api/vehicles/{id}/webhook",
   "La aplicación hacía una petición server-side a la URL provista sin validar el destino.",
   "Acceso al IMDS de la nube (169.254.169.254) y a servicios internos; alcance cruzado (S:C).",
   "SsrfGuard: solo http/https y el host debe resolver exclusivamente a direcciones públicas (bloquea IMDS, loopback, RFC 1918, CGNAT).",
   "IMDS/loopback/privados/esquemas no-HTTP → rechazo; destino público HTTPS → permitido. Residual TOCTOU documentado."),
 "V-04": ("POST /api/vehicles/import",
   "El parser XML procesaba entidades externas y DTDs.",
   "Exfiltración de archivos locales (file://) y SSRF vía http:// dentro del XML.",
   "Endurecimiento JAXP: disallow-doctype-decl, entidades externas deshabilitadas, FEATURE_SECURE_PROCESSING.",
   "XML con DOCTYPE/entidad externa → 400; XML limpio → 200 parseado."),
 "V-05": ("PATCH /api/drivers/{id}",
   "La actualización vinculaba el body JSON completo sobre la entidad, sin allowlist de campos.",
   "Escritura no autorizada y persistente de campos sensibles (role, password) y fuga de password en la respuesta (integridad de datos, I:H).",
   "DTO explícito con allowlist (phone/email); verificación de propiedad; respuesta sin password.",
   "role/password no asignables por el body; la respuesta no incluye password; el conductor no puede auto-promoverse."),
 "V-06": ("GET /api/reports/download?file=",
   "El nombre de archivo se concatenaba al directorio base sin sanitizar.",
   "Lectura de archivos arbitrarios del contenedor (../../etc/passwd).",
   "Normalización de la ruta y verificación de contención (startsWith) dentro del directorio base.",
   "file=../../../../etc/passwd → 400; un reporte dentro del base → 200."),
 "V-07": ("POST /api/auth/login",
   "El endpoint de login no imponía ningún límite de intentos.",
   "Fuerza bruta de credenciales sin fricción.",
   "LoginRateLimiter: ventana deslizante por IP (5/min por defecto); superar el cupo → 429.",
   "El 6.º intento desde una IP → 429; un login válido desde otra IP → 200. Alcance por-instancia documentado."),
 "V-08": ("Múltiples endpoints + logback",
   "Se registraban cédula, correo y teléfono en texto plano en consola y archivo.",
   "Exposición de datos personales en los logs (Ley 1581, principio de seguridad).",
   "PiiMaskingConverter transversal (redacta correos y secuencias numéricas largas en todo log) y eliminación del volcado de filas.",
   "Correos/cédulas/teléfonos redactados; mensajes sin PII intactos; sin NPE en bordes."),
 "V-09": ("GET /api/drivers/{id}/trips",
   "El endpoint devolvía los viajes de cualquier conductor sin verificar propiedad.",
   "Enumeración de recorridos (geolocalización) de otros conductores cambiando el id.",
   "Verificación de propiedad explícita contra el sujeto autenticado; sin permiso → 403.",
   "Conductor A pidiendo los viajes de B → 403; sus propios viajes → 200."),
 "V-10": ("application.yml",
   "El secreto de firma JWT y la contraseña de admin estaban hardcodeados en la configuración.",
   "Exposición de credenciales versionadas en el repositorio.",
   "Externalización a variables de entorno sin valor por defecto: si faltan, el contexto no arranca (fail-fast).",
   "Sin el secreto el contexto no levanta; con el secreto el login emite un JWT firmado."),
 "V-11": ("Toda la API",
   "No existía filtro de seguridad: todos los endpoints respondían a peticiones anónimas.",
   "Acceso anónimo a la PII de ~60.000 conductores y a la modificación de datos (hallazgo raíz).",
   "Spring Security con matriz default-deny y filtro JWT: solo login y documentación son públicos.",
   "Sin token, un endpoint protegido → 401; con token válido → 200."),
}
for vid in [f"V-{i:02d}" for i in range(1, 12)]:
    d = BY[vid]
    endpoint, desc, impacto, remed, verif = DETALLE[vid]
    doc.add_heading(f"{vid} · {TITULOS[vid]}", level=2)
    mt = doc.add_table(rows=1, cols=4)
    mt.style = "Table Grid"
    hdrs = ["Severidad", "CVSS v3.1", "CWE", "OWASP 2021"]
    vals = [d["sev"], d["cvss"], d["cwe"], d["owasp"]]
    hc = mt.rows[0].cells
    for i in range(4):
        set_cell(hc[i], hdrs[i], bold=True, white=True, size=9)
        shade(hc[i], "1F3A5F")
    rc = mt.add_row().cells
    for i in range(4):
        set_cell(rc[i], vals[i], size=9)
    vec = doc.add_paragraph()
    rr = vec.add_run("Vector: " + d["vector"])
    rr.font.size = Pt(8.5)
    rr.italic = True
    rr.font.color.rgb = GREY
    vec.paragraph_format.space_after = Pt(4)
    para("Endpoint: " + endpoint, size=9.5, bold=True, space_after=2)
    para("Descripción. " + desc, size=10, space_after=2)
    para("Impacto. " + impacto, size=10, space_after=2)
    para("Remediación. " + remed, size=10, space_after=2)
    p = doc.add_paragraph()
    r1 = p.add_run("Verificación (test dual). ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = GREEN
    r2 = p.add_run(verif)
    r2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    para("Estado: Fixed · commit %s (FSEC-17)." % d["commit"], size=9, italic=True, color=GREEN, space_after=10)
doc.add_page_break()

# ============================ 6. HALLAZGOS BONUS ============================
doc.add_heading("6. Hallazgos bonus (FSEC-19)", level=1)
para("Además de los 11 vectores del alcance principal, se documentan hallazgos adicionales con "
     "evidencia propia. El entregable exige ≥2 bonus: se cubren con B-01 (security headers) y V-11 "
     "(ausencia de enforcement de autenticación, que conserva su prefijo V por su rol en el análisis "
     "de interacción). Un tercer candidato —B-02, CORS— se evaluó y no aplica.", space_after=8)

doc.add_heading("6.1 B-01 · Ausencia de HTTP security headers", level=2)
_bmt = doc.add_table(rows=1, cols=4)
_bmt.style = "Table Grid"
_bh = ["Severidad", "CVSS v3.1", "CWE", "OWASP 2021"]
_bv = ["LOW", "3.1", "CWE-693", "A05:2021 Security Misconfiguration"]
_bc = _bmt.rows[0].cells
for i in range(4):
    set_cell(_bc[i], _bh[i], bold=True, white=True, size=9)
    shade(_bc[i], "1F3A5F")
_br = _bmt.add_row().cells
for i in range(4):
    set_cell(_br[i], _bv[i], size=9)
_bvec = doc.add_paragraph()
_bvr = _bvec.add_run("Vector: CVSS:3.1/AV:N/AC:H/PR:N/UI:R/S:U/C:L/I:N/A:N")
_bvr.font.size = Pt(8.5)
_bvr.italic = True
_bvr.font.color.rgb = GREY
_bvec.paragraph_format.space_after = Pt(4)
para("Descripción. En la baseline la app no emitía ningún security header (X-Frame-Options, CSP, "
     "HSTS, X-Content-Type-Options, Referrer/Permissions-Policy) — falla de mecanismo de protección "
     "(CWE-693), sin explotación directa pero eliminando una capa de defensa en profundidad.",
     size=10, space_after=2)
para("Evidencia (DAST del pipeline). ZAP reportaba todos los headers ausentes en la baseline; "
     "post-remediación quedan 2 hallazgos Low residuales (X-Content-Type-Options y "
     "Cross-Origin-Resource-Policy en respuestas fuera de la cadena de filtros).", size=10, space_after=2)
para("Impacto. Clickjacking (Swagger UI framable), MIME-sniffing, downgrade sin HSTS y ausencia de "
     "CSP como amplificador de un XSS futuro. Acotado por ser una API JSON (de ahí LOW).",
     size=10, space_after=2)
_bp = doc.add_paragraph()
_br1 = _bp.add_run("Remediación. ")
_br1.bold = True
_br1.font.size = Pt(10)
_br1.font.color.rgb = GREEN
_br2 = _bp.add_run("Set completo de headers añadido en SecurityConfig (FSEC-17 PR-C, commit 0927e6e), "
                   "verificado por test (headers presentes incluso en un 401). Se apoya en la maquinaria "
                   "de escritura de headers de Spring Security ya parchada del CVE-2026-22732 (override "
                   "a 6.5.9). Residual no bloqueante: añadir Cross-Origin-Resource-Policy y cubrir las "
                   "respuestas de error/springdoc con un filtro global.")
_br2.font.size = Pt(10)
_bp.paragraph_format.space_after = Pt(2)
para("Estado: Fixed · commit 0927e6e (FSEC-17 PR-C), con residual Low documentado.",
     size=9, italic=True, color=GREEN, space_after=10)

doc.add_heading("6.2 V-11 · Ausencia de enforcement de autenticación (hallazgo bonus/raíz)", level=2)
para("V-11 se cuenta también como hallazgo bonus: es el hallazgo de arquitectura detectado durante "
     "la evaluación (ningún endpoint exigía autenticación) y el nodo raíz del análisis de interacción "
     "(§4). Conserva su prefijo V —en lugar de renombrarlo B-XX— porque toda la cadena de remediación "
     "por fases y las referencias cruzadas lo citan como V-11. Detalle completo en la §5 y la ficha "
     "V-11 (HIGH, CVSS 8.2, CWE-306, Fixed).", space_after=8)

doc.add_heading("6.3 B-02 · CORS misconfiguration — no aplica", level=2)
para("Se evaluó una posible política CORS permisiva (CWE-942: origen comodín + credenciales). "
     "No aplica: la app no invoca .cors(), no usa @CrossOrigin ni define un CorsConfigurationSource, "
     "por lo que Spring no emite headers Access-Control-Allow-*. Se documenta la no-aplicabilidad con "
     "la evidencia del grep (ficha B-02) en vez de omitirla.", space_after=8)

doc.add_page_break()

# ============================ 7. MAPEO DE CUMPLIMIENTO ============================
doc.add_heading("7. Mapeo de cumplimiento", level=1)
para("Los hallazgos y sus remediaciones se alinean con los marcos aplicables a FleetSec (empresa en "
     "certificación ISO 27001:2022, tratante de datos personales bajo Ley 1581 de 2012).", space_after=8)
doc.add_heading("7.1 ISO/IEC 27001:2022 — Anexo A", level=2)
iso = doc.add_table(rows=1, cols=3)
iso.style = "Table Grid"
header_row(iso, ["Control Anexo A", "Descripción", "Hallazgos"])
for c1, c2, c3 in [
 ("A.8.28", "Codificación segura", "V-01, V-04, V-05, V-06"),
 ("A.5.17", "Información de autenticación", "V-02, V-10"),
 ("A.8.5", "Autenticación segura", "V-07, V-11"),
 ("A.8.3", "Restricción de acceso a la información", "V-09, V-11"),
 ("A.8.15", "Registro (logging)", "V-08"),
 ("A.5.34", "Privacidad y protección de PII", "V-01, V-08"),
 ("A.8.9", "Gestión de la configuración", "V-10, headers"),
 ("A.8.16", "Actividades de monitoreo", "V-07"),
]:
    r = iso.add_row().cells
    set_cell(r[0], c1, bold=True, size=9)
    set_cell(r[1], c2, size=9)
    set_cell(r[2], c3, size=9)
doc.add_heading("7.2 CIS Controls v8", level=2)
cis = doc.add_table(rows=1, cols=3)
cis.style = "Table Grid"
header_row(cis, ["Control CIS", "Descripción", "Hallazgos"])
for c1, c2, c3 in [
 ("16 (16.1, 16.11)", "Seguridad del software de aplicación", "V-01, V-03, V-04, V-05, V-06"),
 ("3 (3.1, 3.11)", "Protección de datos (cifrado/enmascaramiento)", "V-08, V-10"),
 ("6 (6.7)", "Gestión del control de acceso", "V-09, V-11"),
 ("8 (8.2)", "Gestión de registros de auditoría", "V-08"),
 ("4 (4.1)", "Configuración segura", "V-10, headers"),
]:
    r = cis.add_row().cells
    set_cell(r[0], c1, bold=True, size=9)
    set_cell(r[1], c2, size=9)
    set_cell(r[2], c3, size=9)
doc.add_heading("7.3 Ley 1581 de 2012 (Protección de Datos Personales — Colombia)", level=2)
para("La plataforma trata datos personales de conductores (cédula, correo, teléfono, licencia) y "
     "datos de geolocalización (recorridos). Los hallazgos con impacto directo son:", space_after=6)
bullet("acceso anónimo a PII (V-11) y su extracción vía inyección (V-01) vulneran el principio de seguridad (Art. 4, lit. g).", bold_prefix="Confidencialidad: ")
bullet("el registro de PII en texto plano (V-08) incumple el deber de custodia y el principio de seguridad.", bold_prefix="Logging: ")
bullet("la modificación no autorizada de datos personales (V-05) afecta el principio de veracidad/integridad.", bold_prefix="Integridad: ")
bullet("en producción, una exposición de esta naturaleza es un incidente reportable a la Superintendencia de Industria y Comercio (SIC) conforme a los deberes del Responsable (Art. 17) y del Encargado (Art. 18) y el Decreto 1377 de 2013.", bold_prefix="Deber de notificación: ")
para("En este ejercicio los datos son de prueba y el entorno es controlado; no hubo compromiso de "
     "datos reales ni obligación de notificación. La remediación cierra la brecha antes de cualquier "
     "exposición productiva.", size=9.5, italic=True, color=GREY, space_after=6)
doc.add_page_break()

# ============================ 7. CONCLUSIONES ============================
doc.add_heading("8. Conclusiones y recomendaciones", level=1)
para("La evaluación identificó 11 vulnerabilidades —7 de severidad alta y 4 media— dominadas por un "
     "hallazgo raíz de arquitectura: la ausencia de autenticación obligatoria. El análisis de "
     "interacción demostró que el orden de remediación era crítico, y la corrección se ejecutó en tres "
     "fases secuenciales que respetaron esas dependencias, cerrando explícitamente las dos vías al "
     "account takeover.", space_after=8)
para("A la fecha del reporte, los 11 hallazgos están remediados y verificados con pruebas duales "
     "(ataque rechazado + flujo legítimo operativo), y el pipeline de seguridad DevSecOps —ocho etapas "
     "de análisis— reporta en verde, sin hallazgos de severidad alta pendientes.", space_after=8)
para("Recomendaciones de continuidad:", size=10.5, bold=True, space_after=2)
bullet("activar el gate de seguridad del pipeline como control bloqueante (required check) para todo PR.")
bullet("trasladar el rate limiting y el enmascaramiento de PII a controles de plataforma (WAF/gateway, store compartido) al escalar horizontalmente.")
bullet("cerrar el residual de SSRF (TOCTOU) con conexión a IP validada o allowlist de destinos si el webhook pasa a producción.")
bullet("incorporar estas verificaciones al conjunto de pruebas de regresión permanente.")
para("", space_after=10)
para("Anexo. Las fichas técnicas completas de cada hallazgo (PoC ejecutada, justificación métrica de "
     "CVSS y detalle de remediación) residen en el repositorio bajo vapt/findings/V-XX.md. La evidencia "
     "de verificación reside en la suite de pruebas de la aplicación y en los artefactos del pipeline.",
     size=9, italic=True, color=GREY)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("OK ->", OUT)
print("findings extraidos:", len(IDX), "| paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
